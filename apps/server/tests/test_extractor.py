"""链接/文档解析 · 内容提取服务单测（EXT-8/9/10）。

网页/文档路径 mock 掉网络与第三方库，聚焦提取器的判定逻辑与错误处理；
字幕解析用纯函数直接测。
"""
import asyncio
import io

import pytest

from app.services import extractor
from app.services.extractor import ExtractorError


# ---------- 视频/网页 URL 识别 ----------

def test_is_video_url():
    assert extractor.is_video_url("https://www.bilibili.com/video/BV1GJ411x7h7")
    assert extractor.is_video_url("https://b23.tv/abc123")
    assert extractor.is_video_url("https://v.douyin.com/AbCdEf/")
    assert extractor.is_video_url("https://www.kuaishou.com/short-video/xyz")
    assert not extractor.is_video_url("https://www.xiachufang.com/recipe/1005/")
    assert not extractor.is_video_url("https://mp.weixin.qq.com/s/abc")


# ---------- 网页正文提取 ----------

def _html_ok():
    return (
        "<html><head><title>红烧肉家常做法</title></head><body>"
        "<article><h1>红烧肉</h1>"
        "<p>五花肉切块冷水下锅，加姜片料酒焯透，捞出冲净。</p>"
        "<p>干锅小火把肉煸出猪油，冰糖小火炒至琥珀色下肉上色。</p>"
        "<p>加热水没过肉面，加生抽老抽八角桂皮，小火焖40分钟大火收汁。</p>"
        "<p>全程不要开大火，肉才不会柴；糖色要小火，炒苦就前功尽弃。</p>"
        "</article></body></html>"
    )


async def _fake_get_ok(self, url, **kwargs):
    class _Resp:
        status_code = 200
        text = _html_ok()
    return _Resp()


def test_extract_web_ok(monkeypatch):
    monkeypatch.setattr(extractor.httpx.AsyncClient, "get", _fake_get_ok)
    monkeypatch.setattr(
        extractor.trafilatura,
        "extract_metadata",
        lambda html: type("M", (), {"title": "红烧肉家常做法"})(),
    )
    monkeypatch.setattr(
        extractor.trafilatura,
        "extract",
        lambda html: "五花肉切块冷水下锅，加姜片料酒焯透，捞出冲净备用。干锅小火把肉块煸出猪油，倒出大部分猪油。"
        "冰糖下锅小火慢慢炒，炒到琥珀色冒小泡，倒入肉块翻炒均匀上色，加生抽老抽和热水没过肉面，放八角桂皮姜片，"
        "小火焖四十分钟，最后开大火收汁，撒上葱花即可出锅装盘。",
    )
    result = asyncio.run(extractor.extract_web("https://example.com/recipe"))
    assert result["title"] == "红烧肉家常做法"
    assert "焯" in result["text"] and "上色" in result["text"]


async def _fake_get_blocked(self, url, **kwargs):
    class _Resp:
        status_code = 200
        text = "<html><body>请滑动完成验证</body></html>"
    return _Resp()


def test_extract_web_anti_bot_raises(monkeypatch):
    monkeypatch.setattr(extractor.httpx.AsyncClient, "get", _fake_get_blocked)
    monkeypatch.setattr(extractor.trafilatura, "extract", lambda html: "请滑动完成验证")
    with pytest.raises(ExtractorError, match="防爬"):
        asyncio.run(extractor.extract_web("https://example.com/blocked"))


async def _fake_get_404(self, url, **kwargs):
    class _Resp:
        status_code = 404
        text = ""
    return _Resp()


def test_extract_web_http_error(monkeypatch):
    monkeypatch.setattr(extractor.httpx.AsyncClient, "get", _fake_get_404)
    with pytest.raises(ExtractorError, match="404"):
        asyncio.run(extractor.extract_web("https://example.com/missing"))


def test_caption_vtt_to_text():
    raw = (
        "WEBVTT\n\n"
        "00:00:01.000 --> 00:00:03.000\n<c>五花肉</c>切块\n\n"
        "00:00:03.000 --> 00:00:06.000\n冷水下锅焯水\n"
    )
    lines = extractor._caption_to_text(raw, "vtt")
    joined = "\n".join(lines)
    assert "五花肉" in joined and "焯水" in joined
    assert "-->" not in joined


def test_caption_srt_to_text():
    raw = "1\n00:00:01,000 --> 00:00:03,000\n冰糖小火炒糖色\n\n2\n00:00:03,000 --> 00:00:06,000\n下肉上色\n"
    lines = extractor._caption_to_text(raw, "srt")
    assert "糖色" in "\n".join(lines)


# ---------- 文档提取 ----------

def _make_pdf() -> bytes:
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("红烧肉做法")
    doc.add_paragraph("五花肉切块焯水后煸油，炒糖色，加热水焖40分钟。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_pdf_ok(monkeypatch):
    # pypdf 空页 extract_text 为空 → 判"无文本"失败路径；构造足够文本用 mock
    from pypdf import PdfReader

    class _FakeReader:
        pages = [type("P", (), {"extract_text": lambda self: "五花肉焯水。炒糖色。焖40分钟收汁。"})()]

    monkeypatch.setattr(
        extractor,
        "_extract_pdf",
        lambda content: "五花肉切块冷水下锅，加姜片料酒焯透，捞出冲净。干锅小火煸出猪油，冰糖小火炒出糖色再下肉上色，"
        "加热水没过肉面，小火焖四十分钟大火收汁。",
    )
    result = asyncio.run(extractor.extract_document("红烧肉.pdf", _make_pdf()))
    assert result["title"] == "红烧肉.pdf"
    assert "糖色" in result["text"]


def test_extract_docx_ok(monkeypatch):
    monkeypatch.setattr(
        extractor,
        "_extract_docx",
        lambda content: "红烧肉做法\n五花肉切块冷水下锅，加姜片料酒焯透，捞出冲净。干锅小火煸出猪油，冰糖小火炒出糖色"
        "再下肉上色，加热水没过肉面，小火焖四十分钟大火收汁。",
    )
    result = asyncio.run(extractor.extract_document("菜谱.docx", _make_docx()))
    assert "红烧肉" in result["text"]


def test_extract_txt_ok():
    result = asyncio.run(
        extractor.extract_document(
            "菜谱.txt",
            "五花肉切块冷水下锅，加姜片料酒焯透，捞出冲净。干锅小火煸出猪油，冰糖小火炒出糖色再下肉上色，"
            "加热水没过肉面，小火焖四十分钟大火收汁。".encode("utf-8"),
        )
    )
    assert "糖色" in result["text"]


def test_extract_unsupported_extension():
    with pytest.raises(ExtractorError, match="不支持"):
        asyncio.run(extractor.extract_document("菜谱.xlsx", b"x" * 100))


def test_extract_doc_too_short():
    with pytest.raises(ExtractorError, match="没有足够"):
        asyncio.run(extractor.extract_document("菜谱.txt", "好".encode("utf-8")))
