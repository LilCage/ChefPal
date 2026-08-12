"""链接/文档内容提取服务：网页正文、视频字幕/音轨转写、PDF/Word 文本。

对应需求 EXT-8/9/10（视频/网页/文档解析）的内容来源层：
1. 网页：httpx 抓取 + trafilatura 提取正文
2. 视频：yt-dlp 优先提取字幕；无字幕下载音轨 → 百炼 ASR 转写
3. 文档：pypdf（PDF）/ python-docx（Word）/ 纯文本
"""
import asyncio
import html as html_lib
import io
import re
import tempfile
from pathlib import Path

import httpx
import trafilatura

from app.core.config import get_settings

settings = get_settings()


class ExtractorError(Exception):
    """内容提取失败（含网页防爬、视频无字幕且转写失败、文档无文本等）。"""


# 网页抓取：移动端 UA + 常见反爬提示，抓取结果过短/命中反爬视为失败
_WEB_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) AppleWebKit/605.1.15 "
        "(KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9",
}
_WEB_MIN_TEXT = 60
_BLOCK_HINTS = ("请滑动完成验证", "滑动验证", "安全验证", "验证码", "登录后查看", "请先登录", "访问过于频繁")


def is_video_url(url: str) -> bool:
    """按域名判断是否为视频链接（B站/抖音/快手）。"""
    return bool(
        re.search(
            r"(bilibili\.com/video|b23\.tv|douyin\.com|iesdouyin\.com|v\.douyin\.com|kuaishou\.com)",
            url,
            re.I,
        )
    )


# ---------- 网页 ----------

async def extract_web(url: str) -> dict:
    """抓取网页并提取正文。返回 {"title", "text"}；失败抛 ExtractorError。"""
    try:
        async with httpx.AsyncClient(
            headers=_WEB_HEADERS, timeout=20, follow_redirects=True
        ) as client:
            resp = await client.get(url)
    except httpx.HTTPError as exc:
        raise ExtractorError(f"网页访问失败：{exc}") from exc
    if resp.status_code != 200:
        raise ExtractorError(f"网页返回状态码 {resp.status_code}")

    html_text = resp.text
    meta = trafilatura.extract_metadata(html_text) if html_text else None
    title = (meta.title if meta and meta.title else "").strip() or _fallback_title(html_text, url)
    text = (trafilatura.extract(html_text) or "").strip()

    if len(text) < _WEB_MIN_TEXT or any(h in text for h in _BLOCK_HINTS):
        raise ExtractorError(
            "该网页无法读取内容（可能开启防爬验证或需要登录）。"
            "请更换为可访问的菜谱博客 / 公众号文章链接，或改用上传 PDF / Word 文档。"
        )
    return {"title": title, "text": text}


def _fallback_title(html_text: str, url: str) -> str:
    m = re.search(r"<title[^>]*>([^<]+)</title>", html_text, re.I)
    if m:
        return html_lib.unescape(m.group(1)).strip()[:80]
    return url[:80]


# ---------- 视频 ----------

# 优先尝试的中文字幕语言（B站常见）
_SUB_LANGS = ("zh-Hans", "zh-CN", "zh", "zh_CN", "zh-Hant", "zh-TW", "en")


async def extract_video(url: str) -> dict:
    """提取视频字幕/音轨转写。返回 {"title", "text", "source_note"}；失败抛 ExtractorError。

    策略（已与用户确认：字幕优先 + 音轨转写兜底）：
    1) yt-dlp 读取字幕信息 → 抓取字幕文本（无需 ffmpeg）
    2) 无字幕 → 下载最佳音轨（通常 m4a）→ 百炼 ASR 转写
    """
    import yt_dlp

    try:
        with yt_dlp.YoutubeDL({"quiet": True, "no_warnings": True, "noplaylist": True}) as ydl:
            info = await asyncio.to_thread(
                ydl.extract_info, url, download=False
            )
    except Exception as exc:  # noqa: BLE001  yt-dlp 各类下载错误
        raise ExtractorError(f"无法访问该视频：{exc}") from exc

    title = (info.get("title") or "").strip() or url[:60]

    # 1) 字幕优先
    sub_text = await asyncio.to_thread(_pick_subtitle_text, info)
    if sub_text:
        return {"title": title, "text": sub_text, "source_note": "字幕"}

    # 2) 音轨转写兜底
    with tempfile.TemporaryDirectory() as tmp:
        audio_path, audio_mime = await asyncio.to_thread(_download_audio, url, tmp)
        try:
            from app.services import asr as asr_service

            audio_bytes = audio_path.read_bytes()
            text = await asr_service.transcribe_audio(audio_bytes, audio_mime)
        except Exception as exc:  # noqa: BLE001  ASR 失败/格式不支持
            raise ExtractorError(
                f"视频没有字幕，音轨转写也未成功：{exc}。请换一个有字幕的视频，或改传文档。"
            ) from exc
        finally:
            audio_path.unlink(missing_ok=True)

    if not text:
        raise ExtractorError("视频无可用字幕，音轨也未识别到内容")
    return {"title": title, "text": text, "source_note": "音轨转写"}


def _pick_subtitle_text(info: dict) -> str:
    """从 yt-dlp info 中挑选中文字幕并解析成纯文本；无字幕返回空串。"""
    manual = info.get("subtitles") or {}
    auto = info.get("automatic_captions") or {}
    for src in (manual, auto):
        for lang in _SUB_LANGS:
            entries = src.get(lang)
            if not entries:
                continue
            # 优先 vtt/srt 字幕文件
            best = None
            for e in entries:
                if e.get("ext") in ("vtt", "srt"):
                    best = e
                    break
            if best is None and entries:
                best = entries[0]
            if best and best.get("url"):
                try:
                    return _fetch_and_parse_caption(best["url"], best.get("ext") or "vtt")
                except Exception:  # noqa: BLE001
                    continue
    return ""


def _fetch_and_parse_caption(url: str, ext: str) -> str:
    resp = httpx.get(url, timeout=15, headers=_WEB_HEADERS)
    resp.raise_for_status()
    raw = resp.text
    lines = _caption_to_text(raw, ext)
    if len(lines) < 10:
        return ""
    return "\n".join(lines)


def _caption_to_text(raw: str, ext: str) -> list[str]:
    """把 vtt/srt 字幕转成纯文本行（去掉时间轴/序号/HTML 标签）。"""
    if ext == "srt":
        # 去掉序号行与时间轴行
        keep = []
        for line in raw.splitlines():
            s = line.strip()
            if not s or s.isdigit() or "-->" in s:
                continue
            keep.append(s)
    else:  # vtt
        keep = []
        in_header = True
        for line in raw.splitlines():
            s = line.strip()
            if in_header:
                if s == "WEBVTT":
                    continue
                if "-->" not in s and s:
                    in_header = False
            if "-->" in s:
                continue
            if not s:
                continue
            keep.append(s)
    # 去标签 + 去重复字幕块（vtt 常重复两遍）
    cleaned: list[str] = []
    for s in keep:
        s = re.sub(r"<[^>]+>", "", s)
        s = html_lib.unescape(s).strip()
        if s and (not cleaned or cleaned[-1] != s):
            cleaned.append(s)
    return cleaned


def _download_audio(url: str, tmp_dir: str):
    """下载最佳音轨到临时目录，返回 (path, mime)。"""
    import glob

    import yt_dlp

    outtmpl = str(Path(tmp_dir) / "audio.%(ext)s")
    opts = {
        "format": "bestaudio/best",
        "outtmpl": outtmpl,
        "quiet": True,
        "no_warnings": True,
        "noplaylist": True,
    }
    with yt_dlp.YoutubeDL(opts) as ydl:
        ydl.extract_info(url, download=True)
    files = glob.glob(str(Path(tmp_dir) / "audio.*"))
    if not files:
        raise ExtractorError("音频下载失败")
    path = Path(files[0])
    mime = {
        "m4a": "audio/mp4",
        "mp3": "audio/mpeg",
        "aac": "audio/aac",
        "ogg": "audio/ogg",
        "opus": "audio/ogg",
        "wav": "audio/wav",
        "webm": "audio/webm",
    }.get(path.suffix.lstrip(".").lower(), "audio/mpeg")
    return path, mime


# ---------- 文档 ----------

DOC_MAX_BYTES = 10 * 1024 * 1024  # 10MB


async def extract_document(filename: str, content: bytes) -> dict:
    """按扩展名提取 PDF / Word / 纯文本。返回 {"title", "text"}；失败抛 ExtractorError。"""
    if len(content) > DOC_MAX_BYTES:
        raise ExtractorError("文档过大（超过 10MB），请压缩后再试")
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            text = _extract_pdf(content)
        elif suffix in (".docx", ".docm"):
            text = _extract_docx(content)
        elif suffix in (".txt", ".md", ".markdown"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise ExtractorError("暂不支持该文件类型，请上传 PDF 或 Word(.docx)")
    except ExtractorError:
        raise
    except Exception as exc:  # noqa: BLE001  pypdf/python-docx 解析失败
        raise ExtractorError(f"文档解析失败：{exc}") from exc

    text = text.strip()
    if len(text) < _WEB_MIN_TEXT:
        raise ExtractorError("文档中没有足够的文本内容，无法解析")
    return {"title": filename, "text": text}


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t.strip())
    return "\n".join(parts)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
